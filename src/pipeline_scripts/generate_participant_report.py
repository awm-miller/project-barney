#!/usr/bin/env python3

import sqlite3
import re
import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DATABASE_NAME = "pipeline_database.db"
DOCX_FILENAME = "participant_report.docx"

# Mapping for Arabic month names is no longer needed for date parsing from description
# ARABIC_MONTHS_TO_NUM = { ... }

def get_video_data(conn):
    """Fetches id, video_id, title, published_at (DB), and description from the videos table."""
    cursor = conn.cursor()
    try:
        # Ensure published_at is fetched and is not null if we rely on it
        cursor.execute("SELECT id, video_id, title, published_at, description FROM videos WHERE description IS NOT NULL AND description != '' AND published_at IS NOT NULL")
        videos = cursor.fetchall()
        return videos
    except sqlite3.Error as e:
        print(f"Database error fetching video data: {e}")
        return []

# Removed parse_date_from_description function

def parse_participants_from_description(description_text):
    """
    Parses participant names following a phrase like 'يشارك في الحلقة ... كل من:'.
    """
    participants = []
    if not description_text:
        return participants

    intro_regex = re.compile(r"يشارك في الحلقة[^:۱۲۳۴۵۶۷۸۹0123456789]+?كل من:\s*")
    intro_match = intro_regex.search(description_text)

    if intro_match:
        start_of_names_block = intro_match.end()
        names_block_text = description_text[start_of_names_block:]
        
        potential_names = names_block_text.split('\n')
        
        for name_line in potential_names:
            stripped_name = name_line.strip()
            if not stripped_name: 
                if participants: 
                    break
                else: 
                    continue
            
            if re.fullmatch(r"[\sء-ي]+", stripped_name) and len(stripped_name) > 3:
                participants.append(stripped_name)
            else:
                if participants:
                    break
    return participants

def add_rtl_paragraph(document, text, style=None, bold=False, size=None):
    """Adds a paragraph with RTL formatting."""
    p = document.add_paragraph(text, style=style)
    p_format = p.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_format.right_to_left = True
    if bold or size:
        for run in p.runs:
            if bold:
                run.bold = True
            if size:
                run.font.size = Pt(size)
    return p

def format_db_date(published_at_db_str):
    """Converts db timestamp string (e.g., ISO format) to YYYY-MM-DD for display and sorting."""
    if not published_at_db_str:
        return "Unknown Date", "9999-99-99" # For sorting unknown dates last
    try:
        # Attempt to parse common ISO formats
        if 'T' in published_at_db_str:
            dt_obj = datetime.fromisoformat(published_at_db_str.replace('Z', '+00:00'))
        elif ' ' in published_at_db_str: # e.g. YYYY-MM-DD HH:MM:SS
            dt_obj = datetime.strptime(published_at_db_str.split(' ')[0], '%Y-%m-%d')
        else: # Assume YYYY-MM-DD if no time part
            dt_obj = datetime.strptime(published_at_db_str, '%Y-%m-%d')
        
        display_date = dt_obj.strftime("%Y-%m-%d")
        return display_date, display_date # Use the same for display and sorting
    except ValueError:
        print(f"Warning: Could not parse date from database: {published_at_db_str}")
        return published_at_db_str, "9999-99-99" # Return original on error, sort last

def main():
    print(f"Starting participant report generation...")
    conn = sqlite3.connect(DATABASE_NAME)
    if not conn:
        print(f"Error: Could not connect to database '{DATABASE_NAME}'")
        return

    videos_data = get_video_data(conn)
    if not videos_data:
        print("No video descriptions with publication dates found in the database to process.")
        conn.close()
        return

    print(f"Processing {len(videos_data)} videos with descriptions and publication dates...")

    participants_by_date = {}  # Key: sortable_date_str (from DB), Value: (display_date_str, set_of_participants)
    all_participants = set()

    for video_db_id, video_id_yt, title, published_at_db, description in videos_data:
        display_date_str, sortable_date_str = format_db_date(published_at_db)
        participants = parse_participants_from_description(description)

        if participants: # We only care if participants are found
            print(f"  Video {video_id_yt} (Date: {display_date_str}): Found {len(participants)} participants.")
            if sortable_date_str not in participants_by_date:
                participants_by_date[sortable_date_str] = (display_date_str, set())
            
            current_participants_for_date = participants_by_date[sortable_date_str][1]
            current_participants_for_date.update(participants)
            all_participants.update(participants)
        elif description: # Log if description exists but parsing failed for participants
            if not participants:
                 print(f"  Video {video_id_yt} (Date: {display_date_str}): Could not parse participants from description matching pattern.")

    conn.close()
    print("Database connection closed.")

    if not participants_by_date and not all_participants:
        print("No participant data extracted. DOCX file will not be generated.")
        return

    doc = Document()
    
    title_p = doc.add_heading("تقرير المشاركين في الحلقات", level=0)
    title_p_format = title_p.paragraph_format
    title_p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p_format.right_to_left = True

    add_rtl_paragraph(doc, "المشاركون حسب تاريخ الحلقة", bold=True, size=16)
    
    sorted_dates_keys = sorted(participants_by_date.keys()) # Already YYYY-MM-DD or error string

    if not sorted_dates_keys:
        add_rtl_paragraph(doc, "لم يتم العثور على مشاركين بتواريخ محددة.")
    else:
        for sortable_date_key in sorted_dates_keys:
            display_date, date_participants_set = participants_by_date[sortable_date_key]
            
            # Find a video_id associated with this date for display.
            # This assumes one video_id per date is sufficient if multiple videos fall on the same date.
            # A more robust solution might involve restructuring participants_by_date to store video_ids.
            # For simplicity, we'll find the first video_id encountered for this date.
            video_id_for_date = "N/A"
            for vid_db_id, vid_id_yt, _, pub_at_db, desc in videos_data:
                _, s_date = format_db_date(pub_at_db)
                if s_date == sortable_date_key:
                    video_id_for_date = vid_id_yt
                    break
            
            add_rtl_paragraph(doc, f"تاريخ الحلقة: {display_date} (Video ID: {video_id_for_date})", bold=True, size=14)
            if date_participants_set:
                for participant in sorted(list(date_participants_set)):
                    add_rtl_paragraph(doc, f"- {participant}", size=12)
            else:
                add_rtl_paragraph(doc, "  (لا يوجد مشاركون لهذه الحلقة)", size=12)
            doc.add_paragraph()

    doc.add_page_break()
    add_rtl_paragraph(doc, "القائمة الكاملة لجميع المشاركين", bold=True, size=16)
    
    if not all_participants:
        add_rtl_paragraph(doc, "لم يتم العثور على أي مشاركين.")
    else:
        for participant in sorted(list(all_participants)):
            add_rtl_paragraph(doc, participant, size=12)

    try:
        doc.save(DOCX_FILENAME)
        print(f"Participant report successfully saved to '{DOCX_FILENAME}'")
    except Exception as e:
        print(f"Error saving DOCX file: {e}")

if __name__ == "__main__":
    main() 